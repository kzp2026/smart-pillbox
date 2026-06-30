from __future__ import annotations

import argparse
import json
import subprocess
import sys
from pathlib import Path

import pandas as pd

from common import ensure_output_dir, resolve_latest_output_path, save_workbook


EVALUATION_INDICATORS = [
    "需求匹配度",
    "适老化友好性",
    "功能完整性",
    "结构合理性",
    "操作便利性",
    "材料可行性",
    "工程可行性",
    "成本合理性",
    "可优化性",
]


def read_excel(path: Path) -> pd.DataFrame:
    path = resolve_latest_output_path(path)
    if not path.exists():
        return pd.DataFrame()
    try:
        return pd.read_excel(path)
    except Exception:
        return pd.DataFrame()


def read_text(path: Path) -> str:
    path = resolve_latest_output_path(path)
    if not path.exists():
        return ""
    try:
        return path.read_text(encoding="utf-8")
    except Exception:
        return ""


def ensure_ai_parameters(output_dir: Path, product_name: str) -> None:
    parameter_path = resolve_latest_output_path(output_dir / "AI生成参数表.xlsx")
    if parameter_path.exists():
        return
    script_path = Path(__file__).resolve().parent / "07_generate_ai_parameters.py"
    subprocess.run(
        [sys.executable, str(script_path), "--output-dir", str(output_dir), "--product-name", product_name],
        cwd=Path(__file__).resolve().parents[1],
        check=True,
    )


def calculate_score(indicator: str, ai_df: pd.DataFrame, scheme_text: str) -> int:
    base = 82
    if not ai_df.empty:
        base += min(len(ai_df), 8)
    if scheme_text:
        base += 3

    text = scheme_text + " " + " ".join(ai_df.astype(str).head(8).to_numpy().ravel().tolist()) if not ai_df.empty else scheme_text
    keyword_bonus = {
        "需求匹配度": ["需求", "痛点", "评论", "参数"],
        "适老化友好性": ["老人", "老年", "适老", "照护"],
        "功能完整性": ["功能", "核心", "提醒", "支撑", "防滑"],
        "结构合理性": ["结构", "装配", "支撑", "模块"],
        "操作便利性": ["操作", "交互", "便捷", "简单"],
        "材料可行性": ["材料", "工艺", "塑料", "金属", "软胶"],
        "工程可行性": ["工程", "制造", "安装", "维护"],
        "成本合理性": ["成本", "量产", "合理", "标准"],
        "可优化性": ["优化", "迭代", "评价", "改进"],
    }
    base += sum(2 for keyword in keyword_bonus.get(indicator, []) if keyword in text)
    return max(60, min(base, 96))


def build_evaluation_table(product_name: str, ai_df: pd.DataFrame, scheme_text: str) -> pd.DataFrame:
    suggestions = {
        "需求匹配度": "继续保留评论证据链，优先优化高频痛点对应功能。",
        "适老化友好性": "加强字体、握持、反馈和防误触设计，降低老年用户学习成本。",
        "功能完整性": "检查核心功能、辅助功能和异常状态提示是否形成闭环。",
        "结构合理性": "进一步验证受力路径、装配关系、维护方式和安全冗余。",
        "操作便利性": "减少操作步骤，增加清晰反馈和一眼可懂的交互提示。",
        "材料可行性": "结合使用场景补充防滑、抗菌、防水、耐磨和清洁工艺验证。",
        "工程可行性": "补充零部件标准化、制造工艺、安装方式和可靠性测试方案。",
        "成本合理性": "区分基础版与增强版配置，控制非必要传感器和复杂结构成本。",
        "可优化性": "建立用户反馈复测机制，用评分结果反向更新 AI 生成参数。",
    }
    explanations = {
        "需求匹配度": "评价设计方案是否回应评论数据提取出的核心需求与痛点。",
        "适老化友好性": "评价目标用户在认知、握持、视认、行动辅助方面的友好程度。",
        "功能完整性": "评价核心功能、辅助功能、反馈功能是否完整覆盖使用流程。",
        "结构合理性": "评价功能是否能被清晰、稳定、可维护的结构实现。",
        "操作便利性": "评价用户完成主要任务所需步骤、学习成本和错误恢复难度。",
        "材料可行性": "评价材料与工艺是否适配场景、耐用性、安全性和清洁维护要求。",
        "工程可行性": "评价方案从概念到制造、装配、测试和量产的落地可能性。",
        "成本合理性": "评价功能配置、材料选择和结构复杂度是否符合成本约束。",
        "可优化性": "评价方案是否能通过评价结果继续迭代生成参数和设计方案。",
    }
    return pd.DataFrame(
        [
            {
                "评价指标": indicator,
                "分值": calculate_score(indicator, ai_df, scheme_text),
                "评价说明": explanations[indicator],
                "优化建议": suggestions[indicator],
            }
            for indicator in EVALUATION_INDICATORS
        ]
    )


def save_summary_docx(product_name: str, evaluation_df: pd.DataFrame, output_path: Path) -> None:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except Exception:
        print("未安装 python-docx，跳过 DOCX 输出。")
        return

    average_score = round(float(evaluation_df["分值"].mean()), 1) if not evaluation_df.empty else 0
    top_items = evaluation_df.sort_values("分值", ascending=False).head(3)["评价指标"].tolist() if not evaluation_df.empty else []
    weak_items = evaluation_df.sort_values("分值", ascending=True).head(3)["评价指标"].tolist() if not evaluation_df.empty else []

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    doc.add_heading(f"{product_name}开题报告实验结果摘要", level=0)
    doc.add_paragraph(
        "本系统以用户评论数据为输入，经过评论清洗、关键词提取、情感分析、主题聚类、需求映射和 Neo4j 知识图谱构建，"
        "将需求信息转化为 AI 可识别的结构化生成参数，并进一步形成 Prompt 模板、设计方案、设计图片与方案评价结果。"
    )
    doc.add_heading("一、技术路线", level=1)
    doc.add_paragraph("用户评论数据 → 需求提取 → 知识图谱关系路径 → AI 生成参数 → Prompt 模板 → 设计方案生成 → 方案评价与优化。")
    doc.add_heading("二、实验输出", level=1)
    for item in [
        "需求—功能—结构映射表",
        "AI 生成参数表与 JSON 参数",
        "Prompt 模板",
        "产品设计方案与设计图片",
        "方案评价表",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("三、评价结果摘要", level=1)
    doc.add_paragraph(f"方案综合平均分为 {average_score} 分。优势指标包括：{ '、'.join(top_items) if top_items else '暂无' }。")
    doc.add_paragraph(f"后续优化重点包括：{ '、'.join(weak_items) if weak_items else '暂无' }。")
    doc.add_heading("四、开题报告支撑价值", level=1)
    doc.add_paragraph(
        "该结果可用于说明研究中的实验方案、系统流程和技术可行性：系统不是直接主观编写设计方案，"
        "而是通过真实评论证据建立需求来源，再通过知识图谱关系路径转化为可调用的 AI 生成参数。"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="第九阶段：生成方案评价表与开题报告摘要")
    parser.add_argument("--output-dir", default="output", help="输出目录")
    parser.add_argument("--product-name", default="产品", help="产品名称")
    args = parser.parse_args()

    product_name = args.product_name
    output_dir = ensure_output_dir(args.output_dir)
    ensure_ai_parameters(output_dir, product_name)

    ai_df = read_excel(output_dir / "AI生成参数表.xlsx")
    scheme_text = read_text(output_dir / f"{product_name}产品设计方案.txt")
    evaluation_df = build_evaluation_table(product_name, ai_df, scheme_text)
    save_workbook(output_dir / "方案评价表.xlsx", {"方案评价": evaluation_df})
    save_summary_docx(product_name, evaluation_df, output_dir / "开题报告实验结果摘要.docx")
    (output_dir / "方案评价结果.json").write_text(
        json.dumps(
            {
                "product_type": product_name,
                "average_score": round(float(evaluation_df["分值"].mean()), 1),
                "items": evaluation_df.to_dict(orient="records"),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    print(f"产品名称：{product_name}")
    print(f"评价指标数量：{len(evaluation_df)}")
    print(f"已生成：{output_dir / '方案评价表.xlsx'}")
    print(f"已生成：{output_dir / '开题报告实验结果摘要.docx'}")


if __name__ == "__main__":
    main()
