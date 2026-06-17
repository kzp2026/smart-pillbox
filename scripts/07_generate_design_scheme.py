from __future__ import annotations

import argparse
import os
import subprocess
import sys
from pathlib import Path

import pandas as pd

from common import ensure_output_dir


def read_sheet(path: Path, sheet_name: str) -> pd.DataFrame:
    """读取 Excel Sheet，不存在时返回空表。"""
    if not path.exists():
        return pd.DataFrame()
    try:
        return pd.read_excel(path, sheet_name=sheet_name)
    except Exception:
        return pd.DataFrame()


def ensure_mapping_database(output_dir: Path) -> Path:
    """如果映射数据库不存在，自动运行第五阶段脚本生成。"""
    mapping_path = output_dir / "智能药盒需求功能映射数据库.xlsx"
    if mapping_path.exists():
        return mapping_path

    script_path = Path(__file__).resolve().parent / "05_build_mapping_database.py"
    subprocess.run(
        [sys.executable, str(script_path), "--output-dir", str(output_dir)],
        cwd=Path(__file__).resolve().parents[1],
        check=True,
    )
    return mapping_path


def table_to_markdown(df: pd.DataFrame, columns: list[str], max_rows: int = 8) -> str:
    """把表格摘要转换成 Markdown 文本。"""
    if df.empty:
        return "暂无数据。"
    use_cols = [col for col in columns if col in df.columns]
    rows = []
    for _, row in df.head(max_rows).iterrows():
        parts = [f"{col}：{row.get(col, '')}" for col in use_cols]
        rows.append("- " + "；".join(parts))
    return "\n".join(rows)


def build_offline_scheme(
    req_df: pd.DataFrame,
    func_df: pd.DataFrame,
    struct_df: pd.DataFrame,
    opportunity_df: pd.DataFrame,
    topic_df: pd.DataFrame,
) -> str:
    """基于分析结果离线生成产品设计方案。"""
    top_requirements = req_df.sort_values("重要度", ascending=False).head(6) if "重要度" in req_df.columns else req_df.head(6)
    top_opportunities = opportunity_df.head(8)

    topic_text = table_to_markdown(topic_df, ["主题名称", "评论数", "主题关键词"], 6)
    requirement_text = table_to_markdown(top_requirements, ["需求名称", "来源关键词", "情感倾向", "重要度"], 8)
    function_text = table_to_markdown(func_df, ["功能名称", "功能类别", "功能描述"], 10)
    structure_text = table_to_markdown(struct_df, ["结构名称", "结构类型", "结构描述"], 10)
    opportunity_text = table_to_markdown(top_opportunities, ["设计机会点", "证据关键词", "建议功能", "建议结构"], 10)

    return f"""# 智能药盒产品设计方案

## 一、研究数据基础
本方案基于京东智能药盒用户评论数据，通过评论清洗、TF-IDF 关键词提取、中文情感分析、主题聚类和需求-功能-结构映射生成。该流程将用户评论中的显性评价转化为可用于产品设计决策的需求证据。

## 二、用户评论主题概括
{topic_text}

## 三、核心用户需求
{requirement_text}

综合评论数据可以看出，智能药盒设计的核心矛盾集中在“按时提醒、分药防错、老人易用、远程监护、便携防潮、续航可靠”等方面。其中，老人及家属是主要使用与购买人群，产品不仅要完成药品收纳，还要承担用药管理、风险提醒和亲属协同监护的作用。

## 四、产品设计定位
产品定位为“面向老年慢病人群与家庭照护场景的智能服药管理终端”。设计目标包括：
1. 降低漏服、错服和重复服药风险。
2. 降低老年用户学习成本，保持操作简单、提示明确。
3. 支持家庭成员远程查看服药状态，提升照护效率。
4. 兼顾便携、防潮、续航和家庭环境中的外观接受度。

## 五、功能设计方案
{function_text}

功能层面建议采用“多模态提醒 + 分仓管理 + 远程同步 + 适老化交互”的组合。提醒方式不应只依赖单一声音，可采用声音、灯光、语音和手机消息协同；分药结构应与日期、时段、剂量形成直观对应；远程端应突出服药记录查看、异常提醒和家属确认。

## 六、结构设计方案
{structure_text}

结构层面建议采用模块化药仓、密封盖体、低功耗主控、LED 状态灯、蜂鸣器和无线通信模块。药仓应支持拆洗和独立取放，盖体应具备防潮密封能力，外壳采用圆角与柔和配色，降低医疗器械感，增强家庭日常使用的亲和感。

## 七、关键设计机会点
{opportunity_text}

## 八、交互流程建议
1. 初次使用：扫码或按钮引导完成药仓设置、服药时间设置和家属账号绑定。
2. 日常装药：用户按日期或时段将药品放入对应药仓，系统通过灯光或界面提示确认。
3. 到点提醒：药盒发出声音、灯光或语音提醒，同时向手机端推送信息。
4. 服药确认：打开对应药仓或按确认键后记录服药状态。
5. 异常处理：超时未确认时推送家属端，形成远程照护闭环。

## 九、论文实验中的应用价值
该设计方案可作为“用户评论数据驱动产品设计”的实验输出，用于说明从评论文本到用户需求、从需求到功能、从功能到结构的转化路径。输出的映射数据库和知识图谱文件可进一步用于可视化展示需求关联、功能支撑关系和结构实现逻辑。
"""


def maybe_llm_enhance(base_scheme: str) -> tuple[str, str]:
    """可选使用 OpenAI/DeepSeek 兼容接口增强文案；失败时保留离线方案。"""
    api_key = os.getenv("LLM_API_KEY")
    if not api_key:
        return base_scheme, "离线模板生成"

    try:
        from openai import OpenAI
    except Exception:
        return base_scheme + "\n\n> 注：检测到 LLM_API_KEY，但未安装 openai 包，已使用离线模板生成。", "离线模板生成"

    try:
        base_url = os.getenv("LLM_BASE_URL") or None
        model = os.getenv("LLM_MODEL", "deepseek-chat")
        client = OpenAI(api_key=api_key, base_url=base_url)
        prompt = (
            "请在不编造数据的前提下，把下面的智能药盒产品设计方案润色为研究生论文实验输出风格，"
            "保留章节结构，突出用户评论数据、需求映射和产品设计决策之间的关系。\n\n"
            + base_scheme
        )
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
        )
        content = response.choices[0].message.content
        if content and len(content.strip()) > 200:
            return content, f"大模型增强：{model}"
    except Exception as exc:
        return base_scheme + f"\n\n> 注：大模型增强失败，已使用离线模板生成。错误信息：{exc}", "离线模板生成"

    return base_scheme, "离线模板生成"


def save_docx(text: str, output_path: Path) -> None:
    """把 Markdown 风格文本保存为 Word 文档。"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except Exception:
        print("未安装 python-docx，跳过 DOCX 输出。")
        return

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            paragraph = doc.add_heading(stripped[2:], level=0)
            paragraph.style.font.name = "黑体"
            paragraph.style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        elif stripped.startswith("## "):
            paragraph = doc.add_heading(stripped[3:], level=1)
            paragraph.style.font.name = "黑体"
            paragraph.style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped[:2].isdigit() and stripped[2:3] == ".":
            doc.add_paragraph(stripped, style="List Number")
        else:
            doc.add_paragraph(stripped)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    """第七阶段：生成智能药盒产品设计方案。"""
    parser = argparse.ArgumentParser(description="第七阶段：生成智能药盒产品设计方案")
    parser.add_argument("--output-dir", default="output", help="输出目录")
    args = parser.parse_args()

    output_dir = ensure_output_dir(args.output_dir)
    mapping_path = ensure_mapping_database(output_dir)
    topic_path = output_dir / "BERTopic主题聚类结果.xlsx"

    req_df = read_sheet(mapping_path, "用户需求表")
    func_df = read_sheet(mapping_path, "产品功能表")
    struct_df = read_sheet(mapping_path, "产品结构表")
    opportunity_df = read_sheet(mapping_path, "设计机会点")
    topic_df = read_sheet(topic_path, "主题汇总")

    base_scheme = build_offline_scheme(req_df, func_df, struct_df, opportunity_df, topic_df)
    final_scheme, method = maybe_llm_enhance(base_scheme)

    txt_path = output_dir / "智能药盒产品设计方案.txt"
    docx_path = output_dir / "智能药盒产品设计方案.docx"
    txt_path.write_text(final_scheme, encoding="utf-8")
    save_docx(final_scheme, docx_path)

    print(f"方案生成方式：{method}")
    print(f"已生成：{txt_path}")
    print(f"已生成：{docx_path}")


if __name__ == "__main__":
    main()
